"""
Flask API for processing Human Design Guide templates
Deploy this to Render.com, Railway.app, or any Python hosting service
Then call it from Make.com via HTTP module
"""

from flask import Flask, request, send_file, jsonify
from docx import Document
import tempfile
import os

app = Flask(__name__)

def replace_text_in_paragraph(paragraph, key, value):
    """Replace text in paragraph while preserving formatting"""
    if key in paragraph.text:
        full_text = paragraph.text
        new_text = full_text.replace(key, value)
        
        for run in paragraph.runs:
            run.text = ''
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        return True
    return False

@app.route('/process-guide', methods=['POST'])
def process_guide():
    """
    Process a Human Design guide template
    
    Expects JSON:
    {
        "template_url": "https://drive.google.com/...",  # Direct download link
        "replacements": {
            "{{name}}": "Client Name",
            "{{type}}": "Projector",
            ...
        }
    }
    
    Returns: Completed .docx file
    """
    try:
        data = request.json
        
        # Get template - accept either URL or base64 data
        if 'template_data' in data:
            import base64
            template_bytes = base64.b64decode(data['template_data'])
            template_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            template_file.write(template_bytes)
            template_file.close()
            template_path = template_file.name
        elif 'template_url' in data:
            import requests
            response = requests.get(data['template_url'])
            response.raise_for_status()
            template_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            template_file.write(response.content)
            template_file.close()
            template_path = template_file.name
        else:
            return jsonify({"error": "No template_url or template_data provided"}), 400
        
        # Load document
        doc = Document(template_path)
        
        # Replace placeholders
        replacements = data.get('replacements', {})
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if value:
                    # Properly handle Unicode
                    replace_text_in_paragraph(paragraph, key, str(value))
        
        # Save to temp file
        output_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(output_file.name)
        output_file.close()
        
        # Clean up template
        os.unlink(template_path)
        
        # Return file
        return send_file(
            output_file.name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='completed_guide.docx'
        )
        
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/health', methods=['GET'])
def health():
    return jsonify({"status": "healthy"})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
